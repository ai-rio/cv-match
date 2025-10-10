"""
Text extraction service for CV-Match.

Extracts plain text from PDF and DOCX resume files for AI processing.
"""

import io
import logging
from typing import BinaryIO

import docx
from PyPDF2 import PdfReader

logger = logging.getLogger(__name__)


class TextExtractionError(Exception):
    """Raised when text extraction fails."""

    pass


class TextExtractionService:
    """Service for extracting text from resume files."""

    def __init__(self):
        """Initialize text extraction service."""
        self.supported_formats = [".pdf", ".docx"]

    async def extract_text_from_pdf(self, file_content: bytes | BinaryIO) -> str:
        """
        Extract text from a PDF file.

        Args:
            file_content: PDF file content as bytes or file-like object

        Returns:
            Extracted text as string

        Raises:
            TextExtractionError: If extraction fails
        """
        try:
            # Handle bytes input
            if isinstance(file_content, bytes):
                file_obj: io.BytesIO | BinaryIO = io.BytesIO(file_content)
            else:
                file_obj = file_content

            # Read PDF
            pdf_reader = PdfReader(file_obj)

            # Extract text from all pages
            text_parts = []
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            if not text_parts:
                raise TextExtractionError("Nenhum texto encontrado no arquivo PDF")

            # Join all pages with double newline
            full_text = "\n\n".join(text_parts)

            # Clean up text
            full_text = self._clean_text(full_text)

            logger.info(f"Extracted {len(full_text)} characters from PDF")
            return full_text

        except Exception as e:
            logger.exception(f"Error extracting text from PDF: {str(e)}")
            raise TextExtractionError(f"Falha ao extrair texto do PDF: {str(e)}") from e

    async def extract_text_from_docx(self, file_content: bytes | BinaryIO) -> str:
        """
        Extract text from a DOCX file.

        Args:
            file_content: DOCX file content as bytes or file-like object

        Returns:
            Extracted text as string

        Raises:
            TextExtractionError: If extraction fails
        """
        try:
            # Handle bytes input
            if isinstance(file_content, bytes):
                file_obj: io.BytesIO | BinaryIO = io.BytesIO(file_content)
            else:
                file_obj = file_content

            # Read DOCX
            document = docx.Document(file_obj)

            # Extract text from all paragraphs
            text_parts = []
            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Also extract text from tables
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            if not text_parts:
                raise TextExtractionError("Nenhum texto encontrado no arquivo DOCX")

            # Join all parts with newline
            full_text = "\n".join(text_parts)

            # Clean up text
            full_text = self._clean_text(full_text)

            logger.info(f"Extracted {len(full_text)} characters from DOCX")
            return full_text

        except Exception as e:
            logger.exception(f"Error extracting text from DOCX: {str(e)}")
            raise TextExtractionError(f"Falha ao extrair texto do DOCX: {str(e)}") from e

    async def extract_text(self, file_content: bytes | BinaryIO, file_extension: str) -> str:
        """
        Extract text from a resume file based on its extension.

        Args:
            file_content: File content as bytes or file-like object
            file_extension: File extension (e.g., '.pdf', '.docx')

        Returns:
            Extracted text as string

        Raises:
            TextExtractionError: If file format is unsupported or extraction fails
        """
        file_ext = file_extension.lower()

        if file_ext not in self.supported_formats:
            raise TextExtractionError(
                f"Formato de arquivo não suportado: {file_ext}. "
                f"Formatos suportados: {', '.join(self.supported_formats)}"
            )

        logger.info(f"Extracting text from {file_ext} file")

        if file_ext == ".pdf":
            return await self.extract_text_from_pdf(file_content)
        elif file_ext == ".docx":
            return await self.extract_text_from_docx(file_content)
        else:
            raise TextExtractionError(f"Formato de arquivo não suportado: {file_ext}")

    def _clean_text(self, text: str) -> str:
        """
        Clean extracted text by removing extra whitespace and normalizing.

        Args:
            text: Raw extracted text

        Returns:
            Cleaned text
        """
        # Remove excessive whitespace
        lines = [line.strip() for line in text.split("\n")]

        # Remove empty lines
        lines = [line for line in lines if line]

        # Join lines with single newline
        cleaned = "\n".join(lines)

        # Ensure UTF-8 encoding
        cleaned = cleaned.encode("utf-8", errors="ignore").decode("utf-8")

        return cleaned.strip()

    def validate_text_length(
        self, text: str, min_length: int = 50, max_length: int = 50000
    ) -> bool:
        """
        Validate that extracted text has reasonable length.

        Args:
            text: Extracted text
            min_length: Minimum acceptable length (default: 50 chars)
            max_length: Maximum acceptable length (default: 50,000 chars)

        Returns:
            True if text length is valid

        Raises:
            TextExtractionError: If text length is invalid
        """
        text_length = len(text)

        if text_length < min_length:
            raise TextExtractionError(
                f"O texto do currículo é muito curto ({text_length} caracteres). "
                f"Mínimo necessário: {min_length} caracteres."
            )

        if text_length > max_length:
            raise TextExtractionError(
                f"O texto do currículo é muito longo ({text_length} caracteres). "
                f"Máximo permitido: {max_length} caracteres."
            )

        return True


# Singleton instance
_text_extraction_service: TextExtractionService | None = None


def get_text_extraction_service() -> TextExtractionService:
    """Get or create the singleton TextExtractionService instance."""
    global _text_extraction_service

    if _text_extraction_service is None:
        _text_extraction_service = TextExtractionService()

    return _text_extraction_service


# Convenience function for direct import
async def extract_text(file_content: bytes | BinaryIO, file_extension: str) -> str:
    """
    Extract text from a resume file based on its extension.
    Convenience function for backward compatibility.

    Args:
        file_content: File content as bytes or file-like object
        file_extension: File extension (e.g., '.pdf', '.docx')

    Returns:
        Extracted text as string

    Raises:
        TextExtractionError: If file format is unsupported or extraction fails
    """
    service = get_text_extraction_service()
    return await service.extract_text(file_content, file_extension)
